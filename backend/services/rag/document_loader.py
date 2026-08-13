from __future__ import annotations

import warnings
from pathlib import Path
from typing import Any

import pymupdf
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document

from backend.core.config import PROJECT_ROOT
#from backend.services.asset_service import get_machine
from backend.services.asset_service import get_machine_by_id

SUPPORTED_EXTENSIONS = {
    ".docx",
    ".pdf",
}



DOCUMENT_METADATA = {
    "01_Motor-A_Teknik_Kullanim_ve_Izleme_Kilavuzu": {
        "document_id": "MA-MAN-001",
        "document_type": "technical_manual",
    },
    "02_Motor-A_Alarm_ve_Veri_Kalitesi_Katalogu": {
        "document_id": "MA-ALM-001",
        "document_type": "alarm_catalog",
    },
    "03_Motor-A_Bakim_ve_Ilk_Mudahale_Prosedurleri": {
        "document_id": "MA-MNT-001",
        "document_type": "maintenance_procedure",
    },
    "04_Motor-A_Ariza_Teshis_ve_Yonlendirme_Rehberi": {
        "document_id": "MA-TRB-001",
        "document_type": "troubleshooting_guide",
    },
    "05_Motor-A_Gecmis_Olay_ve_Bakim_Kayitlari": {
        "document_id": "MA-INC-001",
        "document_type": "incident_history",
    },
}



####Klasörün bulunması
def get_knowledge_base_path(
    machine_id: str,
) -> Path:
    """
    assets.yaml üzerinden ilgili makinenin
    knowledge base klasörünü bulur.
    """

    machine = get_machine_by_id(machine_id)

    if machine is None:
        raise ValueError(
            f"Machine not found: {machine_id}"
        )

    knowledge_base_config = machine.get(
        "knowledge_base",
        {},
    )

    relative_path = knowledge_base_config.get("path")

    if not relative_path:
        raise ValueError(
            f"Knowledge base path is not defined for: "
            f"{machine_id}"
        )

    path = PROJECT_ROOT / relative_path

    if not path.exists():
        raise FileNotFoundError(
            f"Knowledge base directory not found: {path}"
        )

    return path



####Word tabloloru ile ilgili yardımcı fonksiyonlar
def _table_to_text(table: Table) -> str:
    """
    Word tablosunu satır bazlı sade metne dönüştürür.
    """

    rows: list[str] = []

    for row in table.rows:
        cells = [
            cell.text.strip()
            for cell in row.cells
        ]

        row_text = " | ".join(cells)

        if row_text.strip(" |"):
            rows.append(row_text)

    return "\n".join(rows)


####Word dosyasında yer alan metinleri çıkarma
def extract_docx_text(
    file_path: Path,
) -> str:
    """
    DOCX içerisindeki paragraf ve tabloları
    dokümandaki sıralarını koruyarak çıkarır.
    """

    doc = DocxDocument(file_path)

    parts: list[str] = []

    for block in doc.iter_inner_content():

        if isinstance(block, Paragraph):
            text = block.text.strip()

            if text:
                parts.append(text)

        elif isinstance(block, Table):
            table_text = _table_to_text(block)

            if table_text:
                parts.append(table_text)

    return "\n\n".join(parts)



####PDF 
def extract_pdf_pages(
    file_path: Path,
) -> list[tuple[int, str]]:
    """
    PDF dosyasını sayfa bazında okur.

    Return:
        [
            (1, "page text"),
            (2, "page text"),
            ...
        ]
    """

    pdf = pymupdf.open(file_path)

    pages: list[tuple[int, str]] = []

    try:
        for page_index, page in enumerate(pdf):

            text = page.get_text(
                "text",
                sort=True,
            ).strip()

            if not text:
                continue

            pages.append(
                (
                    page_index + 1,
                    text,
                )
            )

    finally:
        pdf.close()

    return pages

###Metadata Resolver
def resolve_document_metadata(
    file_path: Path,
) -> dict[str, str]:
    """
    Dosya adına göre dokümanın kimliğini
    ve tipini belirler.
    """

    metadata = DOCUMENT_METADATA.get(
        file_path.stem
    )

    if metadata:
        return metadata.copy()

    return {
        "document_id": file_path.stem,
        "document_type": "unknown",
    }

####Document Loader

def load_machine_documents(
    machine_id: str,
) -> list[Document]:
    """
    Bir makineye ait DOCX ve PDF bilgi kaynaklarını
    LangChain Document nesnelerine dönüştürür.
    """

    machine = get_machine_by_id(machine_id)

    if machine is None:
        raise ValueError(
            f"Machine not found: {machine_id}"
        )

    knowledge_base_path = get_knowledge_base_path(
        machine_id
    )

    documents: list[Document] = []

    files = sorted(
        file_path
        for file_path in knowledge_base_path.iterdir()
        if (
            file_path.is_file()
            and file_path.suffix.lower()
            in SUPPORTED_EXTENSIONS
        )
    )

    for file_path in files:

        document_metadata = resolve_document_metadata(
            file_path
        )

        base_metadata: dict[str, Any] = {
            "plant_id": machine["plant_id"],
            "station_id": machine["station_id"],
            "machine_id": machine_id,
            "asset_type": machine.get("type"),
            "document_id": document_metadata[
                "document_id"
            ],
            "document_type": document_metadata[
                "document_type"
            ],
            "source": file_path.name,
            "source_path": str(file_path),
            "file_type": file_path.suffix.lower(),
        }

        # DOCX
        if file_path.suffix.lower() == ".docx":

            text = extract_docx_text(
                file_path
            )

            if not text:
                warnings.warn(
                    f"No readable text found in DOCX: "
                    f"{file_path.name}"
                )
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        **base_metadata,
                        "page_number": None,
                    },
                )
            )

        # PDF
        elif file_path.suffix.lower() == ".pdf":

            pages = extract_pdf_pages(
                file_path
            )

            if not pages:
                warnings.warn(
                    f"No extractable text found in PDF: "
                    f"{file_path.name}. "
                    "The document may require OCR."
                )
                continue

            for page_number, text in pages:

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            **base_metadata,
                            "page_number": page_number,
                        },
                    )
                )

    return documents
